"""
app/modules/export_engine.py
──────────────────────────────
Module 5 — Export Engine

Responsibility:
    Serialize a clean Pandas DataFrame to the requested file format
    and persist it to the output directory.

Supported formats:
    - Excel (.xlsx) — multi-sheet, styled, auto-column-widths
    - Word  (.docx) — cover page + styled table
    - CSV   (.csv)  — UTF-8 with BOM for Excel compatibility
    - JSON  (.json) — records-oriented with metadata envelope

All exporters write to settings.OUTPUT_DIR/<job_id>.<ext>.
"""

from __future__ import annotations

import json
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from loguru import logger
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from app.core.config import settings
from app.core.exceptions import ExportError, UnsupportedFormatError
from app.models.job import OutputFormat
from app.modules.dataset_builder import Dataset


# ── Export result ──────────────────────────────────────────────────────────────

@dataclass
class ExportResult:
    file_path: str
    file_name: str
    file_size_bytes: int
    output_format: OutputFormat
    rows_written: int
    columns_written: int
    truncated: bool = False
    warnings: list[str] = None

    def __post_init__(self) -> None:
        self.warnings = self.warnings or []


# ── Base exporter ──────────────────────────────────────────────────────────────

class BaseExporter(ABC):
    """Abstract base — all concrete exporters implement  _write()."""

    def export(
        self,
        dataset: Dataset,
        job_id: str,
        url: str,
        instructions: str,
    ) -> ExportResult:
        df       = dataset.dataframe
        warnings = []
        truncated = False

        if len(df) == 0:
            warnings.append("Dataset is empty — exporting placeholder file")

        # Enforce row cap per format
        max_rows = self._max_rows()
        if len(df) > max_rows:
            df = df.head(max_rows)
            truncated = True
            warnings.append(
                f"Dataset truncated from {len(dataset.dataframe)} to {max_rows} rows"
            )

        file_name = self._file_name(job_id)
        file_path = os.path.join(settings.OUTPUT_DIR, file_name)

        try:
            self._write(df, file_path, job_id, url, instructions)
        except Exception as exc:
            raise ExportError(f"Export failed: {exc}", context={"format": self.format.value}) from exc

        file_size = os.path.getsize(file_path)
        logger.info(
            "Export complete",
            format=self.format.value,
            file=file_name,
            rows=len(df),
            size_bytes=file_size,
        )

        return ExportResult(
            file_path=file_path,
            file_name=file_name,
            file_size_bytes=file_size,
            output_format=self.format,
            rows_written=len(df),
            columns_written=len(df.columns),
            truncated=truncated,
            warnings=warnings,
        )

    @property
    @abstractmethod
    def format(self) -> OutputFormat: ...

    @abstractmethod
    def _write(self, df: pd.DataFrame, path: str, job_id: str, url: str, instructions: str) -> None: ...

    @abstractmethod
    def _max_rows(self) -> int: ...

    def _file_name(self, job_id: str) -> str:
        ts   = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        ext  = self.format.value
        return f"scrape_{job_id[:8]}_{ts}.{ext}"


# ── Excel exporter ─────────────────────────────────────────────────────────────

class ExcelExporter(BaseExporter):

    format = OutputFormat.EXCEL

    def _max_rows(self) -> int:
        return settings.MAX_ROWS_EXCEL

    def _write(self, df: pd.DataFrame, path: str, job_id: str, url: str, instructions: str) -> None:
        ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            # Main data sheet
            df.to_excel(writer, index=False, sheet_name="Extracted Data")

            # Metadata sheet
            meta = pd.DataFrame([
                {"Property": "Source URL",         "Value": url},
                {"Property": "Instructions",        "Value": instructions},
                {"Property": "Rows Extracted",      "Value": len(df)},
                {"Property": "Columns",             "Value": ", ".join(df.columns)},
                {"Property": "Generated At",        "Value": ts},
                {"Property": "Job ID",              "Value": job_id},
            ])
            meta.to_excel(writer, index=False, sheet_name="Metadata")

        self._apply_styles(path, df)

    def _apply_styles(self, path: str, df: pd.DataFrame) -> None:
        wb = load_workbook(path)
        ws = wb["Extracted Data"]

        HDR_FILL   = PatternFill("solid", fgColor="0F172A")
        HDR_FONT   = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
        ALT_FILL   = PatternFill("solid", fgColor="F8FAFC")
        BODY_FONT  = Font(name="Calibri", size=10)
        THIN_SIDE  = Side(style="thin", color="CBD5E1")
        BORDER     = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

        # Headers
        for cell in ws[1]:
            cell.fill      = HDR_FILL
            cell.font      = HDR_FONT
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border    = BORDER
        ws.row_dimensions[1].height = 30

        # Data rows
        for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            fill = ALT_FILL if row_idx % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
            for cell in row:
                cell.fill      = fill
                cell.font      = BODY_FONT
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border    = BORDER

        # Auto-size columns
        for col_idx, col_cells in enumerate(ws.columns, start=1):
            max_len = max(
                (len(str(cell.value or "")) for cell in col_cells),
                default=0,
            )
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 4, 10), 60)

        ws.freeze_panes = "A2"

        # Style metadata sheet
        if "Metadata" in wb.sheetnames:
            for cell in wb["Metadata"][1]:
                cell.fill = PatternFill("solid", fgColor="1E293B")
                cell.font = Font(bold=True, color="FFFFFF", name="Calibri")

        wb.save(path)


# ── Word exporter ──────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


class WordExporter(BaseExporter):

    format = OutputFormat.WORD

    def _max_rows(self) -> int:
        return settings.MAX_ROWS_WORD

    def _write(self, df: pd.DataFrame, path: str, job_id: str, url: str, instructions: str) -> None:
        doc = Document()
        self._set_margins(doc)
        self._add_cover(doc, url, instructions, len(df), job_id)
        doc.add_page_break()
        self._add_data_table(doc, df)
        self._add_footer_para(doc, job_id)
        doc.save(path)

    def _set_margins(self, doc: Document) -> None:
        for sec in doc.sections:
            sec.top_margin    = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)

    def _add_cover(self, doc: Document, url: str, instructions: str, rows: int, job_id: str) -> None:
        title = doc.add_paragraph()
        run   = title.add_run("Web Scraping Report")
        run.bold                = True
        run.font.size           = Pt(24)
        run.font.color.rgb      = RGBColor(0x0F, 0x17, 0x2A)
        title.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()
        tbl  = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        labels = ["Source URL", "Instructions", "Records", "Job ID", "Generated At"]
        values = [url, instructions, str(rows), job_id, datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")]

        for i, (lbl, val) in enumerate(zip(labels, values)):
            tbl.rows[i].cells[0].text = lbl
            tbl.rows[i].cells[1].text = val
            _set_cell_shading(tbl.rows[i].cells[0], "1E293B")
            p0 = tbl.rows[i].cells[0].paragraphs[0]
            if p0.runs:
                p0.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p0.runs[0].bold = True

    def _add_data_table(self, doc: Document, df: pd.DataFrame) -> None:
        doc.add_heading("Extracted Data", level=1)
        if df.empty:
            doc.add_paragraph("No data was extracted.")
            return

        cols   = list(df.columns)
        table  = doc.add_table(rows=1 + len(df), cols=len(cols))
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0]
        for i, col in enumerate(cols):
            hdr.cells[i].text = col.replace("_", " ").title()
            _set_cell_shading(hdr.cells[i], "0F172A")
            p = hdr.cells[i].paragraphs[0]
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.runs[0].bold           = True
                p.runs[0].font.size      = Pt(9)

        # Data rows
        for row_idx, (_, record) in enumerate(df.iterrows()):
            row = table.rows[row_idx + 1]
            bg  = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, col in enumerate(cols):
                cell = row.cells[col_idx]
                val  = str(record.get(col, "") or "")[:400]
                cell.text = val
                _set_cell_shading(cell, bg)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)

    def _add_footer_para(self, doc: Document, job_id: str) -> None:
        doc.add_paragraph()
        footer = doc.add_paragraph()
        run = footer.add_run(f"Generated by AI Scraper  •  Job {job_id}  •  {datetime.utcnow().strftime('%Y-%m-%d')}")
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)


# ── CSV exporter ───────────────────────────────────────────────────────────────

class CSVExporter(BaseExporter):

    format = OutputFormat.CSV

    def _max_rows(self) -> int:
        return settings.MAX_ROWS_EXCEL   # Same cap as Excel

    def _write(self, df: pd.DataFrame, path: str, job_id: str, url: str, instructions: str) -> None:
        df.to_csv(path, index=False, encoding="utf-8-sig")   # BOM for Excel compatibility

    def _file_name(self, job_id: str) -> str:
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return f"scrape_{job_id[:8]}_{ts}.csv"


# ── JSON exporter ──────────────────────────────────────────────────────────────

class JSONExporter(BaseExporter):

    format = OutputFormat.JSON

    def _max_rows(self) -> int:
        return settings.MAX_ROWS_EXCEL

    def _write(self, df: pd.DataFrame, path: str, job_id: str, url: str, instructions: str) -> None:
        envelope = {
            "meta": {
                "job_id":       job_id,
                "source_url":   url,
                "instructions": instructions,
                "total_records": len(df),
                "columns":      list(df.columns),
                "generated_at": datetime.utcnow().isoformat(),
            },
            "records": df.where(pd.notna(df), None).to_dict(orient="records"),
        }
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(envelope, fh, indent=2, default=str, ensure_ascii=False)

    def _file_name(self, job_id: str) -> str:
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return f"scrape_{job_id[:8]}_{ts}.json"


# ── Factory ────────────────────────────────────────────────────────────────────

_EXPORTERS: dict[OutputFormat, type[BaseExporter]] = {
    OutputFormat.EXCEL: ExcelExporter,
    OutputFormat.WORD:  WordExporter,
    OutputFormat.CSV:   CSVExporter,
    OutputFormat.JSON:  JSONExporter,
}


class ExportEngine:
    """
    Facade — pick the right exporter based on output format and call it.

    Usage:
        engine = ExportEngine()
        result = engine.export(dataset, job_id, url, instructions, OutputFormat.EXCEL)
    """

    def export(
        self,
        dataset: Dataset,
        job_id: str,
        url: str,
        instructions: str,
        output_format: OutputFormat,
    ) -> ExportResult:
        exporter_cls = _EXPORTERS.get(output_format)
        if exporter_cls is None:
            raise UnsupportedFormatError(
                f"Format '{output_format}' is not supported",
                context={"supported": list(_EXPORTERS.keys())},
            )
        exporter = exporter_cls()
        return exporter.export(dataset, job_id, url, instructions)
